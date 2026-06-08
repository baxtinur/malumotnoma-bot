"""
Маълумотнома Bot
- Botga .docx fayl yuboring
- Bot avtomatik Word+JPG yaratib guruhga yuboradi
"""
import os
import re
import glob
import urllib.request
import urllib.parse
import json
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

TOKEN = "8262234735:AAElRm7XGdKhSzx_wNCJYQs9PfCUS21D4UQ"
GROUP_CHAT_ID = "-1003620751917"
BASE = f"https://api.telegram.org/bot{TOKEN}"

# ─── Telegram yordamchi funksiyalar ────────────────────────────────────────

def tg_get(method, params=None):
    url = f"{BASE}/{method}"
    if params:
        url += "?" + urllib.parse.urlencode(params)
    with urllib.request.urlopen(url, timeout=30) as r:
        return json.loads(r.read())

def tg_send_message(chat_id, text):
    tg_get("sendMessage", {"chat_id": chat_id, "text": text})

def tg_send_file(chat_id, filepath, caption="", method="sendDocument"):
    with open(filepath, "rb") as f:
        file_data = f.read()
    filename = os.path.basename(filepath)
    boundary = "TGBot12345"
    body = b""
    for name, value in [("chat_id", str(chat_id)), ("caption", caption)]:
        body += f"--{boundary}\r\n".encode()
        body += f"Content-Disposition: form-data; name=\"{name}\"\r\n\r\n".encode()
        body += f"{value}\r\n".encode()
    field = "photo" if method == "sendPhoto" else "document"
    body += f"--{boundary}\r\n".encode()
    body += f"Content-Disposition: form-data; name=\"{field}\"; filename=\"{filename}\"\r\n".encode()
    body += b"Content-Type: application/octet-stream\r\n\r\n"
    body += file_data
    body += f"\r\n--{boundary}--\r\n".encode()
    req = urllib.request.Request(
        f"{BASE}/{method}",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"}
    )
    with urllib.request.urlopen(req, timeout=60) as r:
        return json.loads(r.read())

def tg_download_file(file_id, dest_path):
    info = tg_get("getFile", {"file_id": file_id})
    file_path = info["result"]["file_path"]
    url = f"https://api.telegram.org/file/bot{TOKEN}/{file_path}"
    urllib.request.urlretrieve(url, dest_path)

# ─── Converter funksiyalari ─────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_para(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Cambria'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_vmerge(cell, restart=False):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    if restart:
        vMerge.set(qn('w:val'), 'restart')
    tcPr.append(vMerge)

def normalize_category(text):
    t = text.upper().strip()
    if any(x in t for x in ['SIYOS', 'СИЁС']):
        return 'СИЁСИЙ ТАДБИРЛАР'
    if any(x in t for x in ['MADAN', 'МАДАНИЙ', 'MADANIY']):
        return 'МАДАНИЙ ТАДБИРЛАР'
    if any(x in t for x in ['SPORT', 'СПОРТ']):
        return 'СПОРТ ТАДБИРЛАР'
    if any(x in t for x in ['BOSHQ', 'БОШҚ']):
        return 'БОШҚА ТАДБИРЛАР'
    return text

def extract_rows(filepath):
    doc = Document(filepath)
    rows = []
    current_category = None
    main_table = doc.tables[1]
    STOP_WORDS = ['JAMI', 'ЖАМИ']

    for ri, row in enumerate(main_table.rows):
        if ri <= 1:
            continue
        cells = [cell.text.strip() for cell in row.cells]
        first_cell = cells[0]
        full_text = cells[1].strip() if len(cells) > 1 else ''

        if any(s in first_cell.upper() for s in STOP_WORDS):
            break

        unique_cells = set(c.upper() for c in cells if c.strip())
        if len(unique_cells) == 1:
            cat_text = cells[0]
            cat_up = cat_text.upper()
            if any(x in cat_up for x in ['SIYOS', 'MADAN', 'SPORT', 'BOSHQ',
                                          'СИЁС', 'МАДАНИЙ', 'СПОРТ', 'БОШҚ']):
                current_category = normalize_category(cat_text)
                continue

        if not full_text:
            continue

        mg_parts = []
        for ci in range(5, 13):
            val = cells[ci].strip() if len(cells) > ci else ''
            if val and val != '-':
                try:
                    mg_parts.append(int(val.split()[0]))
                except:
                    pass
        total_mg = sum(mg_parts) if mg_parts else 0

        iibb = cells[4].strip() if len(cells) > 4 else ''
        try:
            iibb_n = int(iibb) if iibb and iibb != '-' else 0
        except:
            iibb_n = 0
        jalb_etilgan = str(iibb_n + total_mg) if (iibb_n + total_mg) > 0 else '-'

        rows.append({
            'category': current_category or '',
            'tadbir': full_text,
            'joy': cells[2].strip() if len(cells) > 2 else '',
            'ishtirokchilar': cells[3].strip() if len(cells) > 3 else '',
            'jalb_etilgan': jalb_etilgan,
            'masul': cells[13].strip() if len(cells) > 13 else '',
            'tashabbuskor': cells[14].strip() if len(cells) > 14 else '',
        })
    return rows

def detect_date(filepath):
    basename = os.path.basename(filepath)
    m = re.search(r'(\d{1,2})[_\-\.](\d{2})[_\-\.](\d{4})', basename)
    if m:
        return f"{int(m.group(1)):02d}.{m.group(2)}.{m.group(3)}"
    return ""

def create_malumotnoma(rows, sana, output_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(2)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"2026-йил {sana} куни Тошкент шаҳрида\nўтказиладиган оммавий тадбирлар тўғрисида")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Cambria'

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(3)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("МАЪЛУМОТНОМА")
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.name = 'Cambria'

    col_widths = [0.5, 7.0, 3.3, 0.5, 0.5, 3.5, 3.5]
    headers = [
        'Т.р', 'Тадбирнинг мазмуни, ўтказиладиган\n(бошланадиган) вақти',
        'Тадбир ўтказиладиган жойи',
        'Тадбир иштирокчилар сони',
        'Тадбирга жалб этилган шахсий таркиб',
        'Бириктирилган масъул раҳбар',
        'Тадбир ташаббускори'
    ]

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hrow = table.rows[0]
    for ci, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hrow.cells[ci]
        set_cell_bg(cell, 'D9E1F2')
        set_col_width(cell, w)
        set_cell_margins(cell)
        cell_para(cell, hdr, bold=True, size=10.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    current_cat = None
    counter = 1
    total_ishtirokchi = 0
    total_jalb = 0

    for row_data in rows:
        cat = row_data['category']
        if cat != current_cat:
            current_cat = cat
            cat_row = table.add_row()
            cat_cell = cat_row.cells[0].merge(cat_row.cells[6])
            set_cell_bg(cat_cell, 'BDD7EE')
            set_cell_margins(cat_cell)
            cell_para(cat_cell, cat, bold=True, size=10.5)
            cat_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        drow = table.add_row()
        vals = [str(counter), row_data['tadbir'], row_data['joy'],
                row_data['ishtirokchilar'], row_data['jalb_etilgan'],
                row_data['masul'], row_data['tashabbuskor']]

        for ci, (val, w) in enumerate(zip(vals, col_widths)):
            c = drow.cells[ci]
            set_col_width(c, w)
            set_cell_margins(c)
            if ci == 1:
                c.text = ''
                para = c.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lines = [l.strip() for l in val.split('\n') if l.strip()]
                past_paren = False
                for li, line in enumerate(lines):
                    if li > 0:
                        br = para.add_run()
                        br.add_break()
                        br.font.size = Pt(10.5)
                    run = para.add_run(line)
                    run.font.name = 'Cambria'
                    run.font.size = Pt(10.5)
                    run.bold = past_paren
                    if ')' in line:
                        past_paren = True
            elif ci in [3, 4]:
                cell_para(c, val, bold=True, size=10.5)
            else:
                cell_para(c, val, bold=False, size=10.5)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        try:
            total_ishtirokchi += int(row_data['ishtirokchilar'].replace(' ', ''))
        except:
            pass
        try:
            v = row_data['jalb_etilgan']
            if v != '-':
                total_jalb += int(v.replace(' ', ''))
        except:
            pass
        counter += 1

    # ЖАМИ qatori
    jami_row = table.add_row()
    jami_cell = jami_row.cells[0].merge(jami_row.cells[2])
    set_cell_bg(jami_cell, 'D9E1F2'); set_cell_margins(jami_cell)
    cell_para(jami_cell, 'ЖАМИ:', bold=True, size=10.5)
    jami_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ci, val in enumerate([str(total_ishtirokchi), str(total_jalb) if total_jalb else '', '', '']):
        cx = jami_row.cells[ci + 3]
        set_cell_bg(cx, 'D9E1F2'); set_cell_margins(cx)
        cell_para(cx, val, bold=True, size=10.5)
        cx.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Toifalar yig'indisi
    cat_counts = {}
    for r in rows:
        cat_counts[r['category']] = cat_counts.get(r['category'], 0) + 1
    total_count = sum(cat_counts.values())

    cat_display = [
        ('Сиёсий тадбирлар',  'СИЁСИЙ ТАДБИРЛАР'),
        ('Маданий тадбирлар', 'МАДАНИЙ ТАДБИРЛАР'),
        ('Спорт тадбирлари',  'СПОРТ ТАДБИРЛАР'),
        ('Бошқа тадбирлар',   'БОШҚА ТАДБИРЛАР'),
    ]

    def set_vmerge(cell, restart=False):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vMerge = OxmlElement('w:vMerge')
        if restart:
            vMerge.set(qn('w:val'), 'restart')
        tcPr.append(vMerge)

    for i, (display_name, key) in enumerate(cat_display):
        srow = table.add_row()
        left = srow.cells[0].merge(srow.cells[2])
        set_cell_bg(left, 'BDD7EE'); set_cell_margins(left)
        set_vmerge(left, restart=(i == 0))
        if i == 0:
            cell_para(left, f'Жами тадбирлар: {total_count}', bold=True, size=10.5)
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        mid = srow.cells[3].merge(srow.cells[5])
        set_cell_bg(mid, 'BDD7EE'); set_cell_margins(mid)
        cell_para(mid, display_name, bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        mid.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right = srow.cells[6]
        set_cell_bg(right, 'BDD7EE'); set_cell_margins(right)
        cell_para(right, str(cat_counts.get(key, 0)), bold=True, size=10.5)
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Imzo
    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sr = sign.add_run("Тошкент шаҳар ИИББ ЖТСБ")
    sr.bold = True; sr.font.size = Pt(10.5); sr.font.name = 'Cambria'

    doc.save(output_path)
    return output_path

def docx_to_jpg(docx_path, jpg_path):
    """Word ni JPG ga aylantirish uchun python-docx2pdf yo'q, shuning uchun skip"""
    return None

# ─── Bot asosiy sikl ────────────────────────────────────────────────────────

def process_docx(file_id, filename, user_chat_id):
    tmp_dir = tempfile.mkdtemp()
    input_path = os.path.join(tmp_dir, filename)
    
    tg_send_message(user_chat_id, "⏳ Fayl qabul qilindi, qayta ishlanmoqda...")
    
    # Faylni yuklab olish
    tg_download_file(file_id, input_path)
    
    # Sana aniqlash
    sana = detect_date(input_path)
    sana_safe = sana.replace('.', '_') if sana else 'malumotnoma'
    
    # Maʼlumot ajratish
    rows = extract_rows(input_path)
    
    if not rows:
        tg_send_message(user_chat_id, "❌ Fayldan ma'lumot topilmadi. To'g'ri Кунлик Режа faylini yuboring.")
        return
    
    # Word yaratish
    docx_out = os.path.join(tmp_dir, f"Маълумотнома_{sana_safe}.docx")
    create_malumotnoma(rows, sana, docx_out)
    
    caption = f"📋 {sana} kundagi оммавий тадбирлар маълумотномаси\n✅ {len(rows)} ta tadbir"
    
    # Guruhga Word yuborish
    tg_send_file(GROUP_CHAT_ID, docx_out, caption=caption, method="sendDocument")
    
    # Foydalanuvchiga xabar
    tg_send_message(user_chat_id, f"✅ Tayyor! Guruhga yuborildi.\n📊 {len(rows)} ta tadbir topildi.")

def run_bot():
    print("🤖 Bot ishga tushdi...")
    last_update_id = None
    
    while True:
        try:
            params = {"timeout": 30, "allowed_updates": ["message"]}
            if last_update_id:
                params["offset"] = last_update_id + 1
            
            data = tg_get("getUpdates", params)
            updates = data.get("result", [])
            
            for update in updates:
                last_update_id = update["update_id"]
                msg = update.get("message", {})
                chat_id = msg.get("chat", {}).get("id")
                
                # .docx fayl keldi
                doc = msg.get("document")
                if doc and doc.get("file_name", "").endswith(".docx"):
                    process_docx(doc["file_id"], doc["file_name"], chat_id)
                
                # /start buyrug'i
                elif msg.get("text", "").startswith("/start"):
                    tg_send_message(chat_id, 
                        "Salom! 👋\n\nMen Маълумотнома botiman.\n\n"
                        "📎 Кунлик Режа .docx faylini yuboring — "
                        "men avtomatik Маълумотнома yaratib guruhga yuboraman!")
        
        except KeyboardInterrupt:
            print("Bot to'xtatildi.")
            break
        except Exception as e:
            print(f"Xato: {e}")
            import time; time.sleep(5)

if __name__ == "__main__":
    run_bot()
